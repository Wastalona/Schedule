import re
from os import remove
from datetime import datetime, date

import wget
import requests
from bs4 import BeautifulSoup
from docx import Document

from permanent_schedule import giveThisDay


def downloadDocxFile():
	#Const
	url = 'http://www.bobruisk.belstu.by/dnevnoe-otdelenie/raspisanie-zanyatiy-i-zvonkov-zamenyi'
	response = requests.get(url)
	soup = BeautifulSoup(response.text, 'lxml')

	#Search
	quotes = soup.find_all('address', class_='la_download_link la_preview_1 la_ext_docx')
	quotes_ = soup.find_all('small', class_='la_f_block la_element_footer_date')

	#Filters
	dateOfUpdate = re.findall(r'[А-я]+ \d+ [А-я]+ \d+', str(quotes_[1]))[0]
	address = re.findall(r'(?<=["\'])[^"\']+', str(quotes[0]))[2]

	if int(re.findall(r'\d\d', dateOfUpdate)[0]) == int(datetime.now().day):
		#Download docx file
		wget.download((f'http://www.bobruisk.belstu.by{address}/schedule.docx'))
	else:
		return 0

def findReplacement():
	global date_
	wordDoc = Document("Zamena_SAYT.docx")
	date_ = (
		re.findall(r'\d+ [а-я]+ \d+', wordDoc.paragraphs[0].text)[0],\
		re.findall(r'\(([а-я]+)\)', wordDoc.paragraphs[0].text)[0])
	table_size_row = len(wordDoc.tables[0].rows)
	table_size_col = len(wordDoc.tables[0].columns)
	report, dataframe = {}, []

	# all lines
	for i in range(1,table_size_row):
	    dataframe.append(wordDoc.tables[0].rows[i].cells[0].text)

	# iterate over rows
	for i in dataframe:
		if i == 'ПО 5':
			idGroup = dataframe.index(i)

			# end of replacement
			pattern = dataframe.index(re.findall(r'[А-Я]+ \d+', \
				''.join(dataframe[idGroup+1:]))[0])

			# span indices
			deck = []
			for id in range(len(dataframe[idGroup:pattern])):
				deck.append(idGroup+id)

			# dictionary with substitutions
			out, c = [], 0
			for index in deck:
				for m in range(3, table_size_col, 2):
					out.append(wordDoc.tables[0].rows[index+1].cells[m].text)
				report[wordDoc.tables[0].rows[index+1].cells[1].text] = out[c]#(out[c], out[c+1])
				c += 2
			# print((f'\n{report}'))
			return report
			break


def destruction(): remove('Zamena_SAYT.docx')


def createNewSchedule(zamena, old_sc):
	new_schedule = old_sc.copy()
	for key in zamena.keys():
		if key == re.findall(r'\d/\d', key):
			new_schedule[key] = zamena[key]
		else:
			for i in range(1,3):
				new_schedule[f'{key}/{str(i)}'] = zamena[key]

	return new_schedule


def getResult(item):
	result = ""
	for i, j in item.items():
		result += f'{i} | {j}\n'
	return result[:len(result)-1]


def main():
	if downloadDocxFile() == False: return 'Замен нет'
	else:
		zamena = findReplacement()
		destruction()
		return getResult(createNewSchedule(zamena, giveThisDay(date_[1])))


if __name__ == '__main__':
	main()
